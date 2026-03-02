"""
Create a Word document template (.docx) with pre-configured page layout
matching the Becky story PDF spec:

- Page size: 5.5 × 8.5 inches
- Mirror margins: inside 1in (0.5 + 0.5 gutter), outside 0.5in, top 0.75in, bottom 0.5in
- Page numbers: top of page, outside edge (alternating odd/even)
- Font: Calibri 14pt
- Cover page section (no header/footer) + blank verso
- Copyright: © Thomas Butler. All rights reserved.

The MCP word-document-server can then copy this template and populate it with content.
"""

from docx import Document
from docx.shared import Inches, Pt, Emu, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy


def create_template(output_path: str):
    doc = Document()

    # ── SECTION 1: Cover page ──────────────────────────────────
    cover_section = doc.sections[0]
    _set_page_size(cover_section)
    cover_section.left_margin = Inches(0.75)
    cover_section.right_margin = Inches(0.75)
    cover_section.top_margin = Inches(1.0)
    cover_section.bottom_margin = Inches(1.0)

    # No header/footer on cover
    cover_section.different_first_page_header_footer = False
    _suppress_header_footer(cover_section)

    # Cover content — all paragraphs suppress space_after to control layout
    # Usable height ≈ 468pt (8.5in - 2in margins). Layout: ~1/5 top, title+sub, ~3/5, author+copy, ~1/5 bottom

    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(80)
    p_spacer.paragraph_format.space_after = Pt(0)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(0)
    run = p_title.add_run("STORY TITLE")
    run.font.name = "Calibri"
    run.font.size = Pt(42)
    run.bold = True

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(6)
    p_sub.paragraph_format.space_after = Pt(0)
    run = p_sub.add_run("A Subtitle")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.italic = True

    # Author — push down but not enough to overflow
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.paragraph_format.space_before = Pt(150)
    p_author.paragraph_format.space_after = Pt(0)
    run = p_author.add_run("Thomas Butler")
    run.font.name = "Calibri"
    run.font.size = Pt(13)

    # Copyright
    p_copy = doc.add_paragraph()
    p_copy.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_copy.paragraph_format.space_before = Pt(8)
    p_copy.paragraph_format.space_after = Pt(0)
    run = p_copy.add_run("\u00a9 Thomas Butler. All rights reserved.")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(120, 120, 120)

    # ── Page break → blank verso (back of cover) ──────────────
    doc.add_page_break()
    p_blank = doc.add_paragraph()  # intentionally blank page
    p_blank.paragraph_format.space_before = Pt(0)
    p_blank.paragraph_format.space_after = Pt(0)

    # ── SECTION 2: Story pages ─────────────────────────────────
    # Add a section break (new page) for the story content
    doc.add_section()
    story_section = doc.sections[1]
    _set_page_size(story_section)

    # Mirror margins: inside gets the gutter
    story_section.left_margin = Inches(1.0)    # inside (gutter side) for odd pages
    story_section.right_margin = Inches(0.5)   # outside
    story_section.top_margin = Inches(0.75)
    story_section.bottom_margin = Inches(0.5)
    story_section.gutter = Inches(0)  # handled via mirror margins

    # Enable mirror margins via XML (python-docx doesn't expose this directly)
    sectPr = story_section._sectPr
    mirrorMargins = parse_xml(f'<w:mirrorMargins {nsdecls("w")} />')
    # Set mirror margins at document level (settings.xml)
    _enable_mirror_margins(doc)

    # Enable different odd/even headers
    _enable_odd_even_headers(doc)

    # Set up headers with page numbers
    _setup_page_number_headers(story_section)

    # Set default font for the document
    style = doc.styles['Normal']
    style.font.name = style.font.name or "Calibri"
    style.font.name = "Calibri"
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.4
    style.paragraph_format.space_after = Pt(10)

    # Add a placeholder paragraph so the story section isn't empty
    p_placeholder = doc.add_paragraph()
    run = p_placeholder.add_run("[Story content goes here — delete this paragraph and use MCP tools to add content]")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(180, 180, 180)

    doc.save(output_path)
    print(f"Template created: {output_path}")


def _set_page_size(section):
    """Set page size to 5.5 × 8.5 inches."""
    section.page_width = Inches(5.5)
    section.page_height = Inches(8.5)
    section.orientation = WD_ORIENT.PORTRAIT


def _suppress_header_footer(section):
    """Remove header and footer from a section."""
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.text = ""
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.text = ""


def _enable_mirror_margins(doc):
    """Enable mirror margins at the document settings level."""
    settings = doc.settings.element
    existing = settings.findall(qn('w:mirrorMargins'))
    for e in existing:
        settings.remove(e)
    mirror_elem = parse_xml(f'<w:mirrorMargins {nsdecls("w")} />')
    settings.append(mirror_elem)


def _enable_odd_even_headers(doc):
    """Enable different odd/even headers at the document settings level."""
    settings = doc.settings.element
    existing = settings.findall(qn('w:evenAndOddHeaders'))
    for e in existing:
        settings.remove(e)
    elem = parse_xml(f'<w:evenAndOddHeaders {nsdecls("w")} />')
    settings.append(elem)


def _setup_page_number_headers(section):
    """Set up headers with page numbers on the outside edge.

    Odd pages (recto): page number aligned right
    Even pages (verso): page number aligned left
    """
    sectPr = section._sectPr

    # -- Odd (default) header: page number right-aligned --
    odd_header = section.header
    odd_header.is_linked_to_previous = False
    _clear_header(odd_header)
    p = odd_header.paragraphs[0] if odd_header.paragraphs else odd_header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_page_number_field(p, size=10, color=RGBColor(100, 100, 100))

    # -- Even header: page number left-aligned --
    even_header = section.even_page_header
    even_header.is_linked_to_previous = False
    _clear_header(even_header)
    p = even_header.paragraphs[0] if even_header.paragraphs else even_header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_page_number_field(p, size=10, color=RGBColor(100, 100, 100))


def _clear_header(header):
    """Remove all content from a header."""
    for p in header.paragraphs:
        for run in p.runs:
            run.text = ""


def _add_page_number_field(paragraph, size=10, color=None):
    """Add a PAGE field code to a paragraph (renders as the page number)."""
    run = paragraph.add_run()
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color

    # Insert PAGE field via XML
    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar_begin)

    run2 = paragraph.add_run()
    run2.font.name = "Calibri"
    run2.font.size = Pt(size)
    if color:
        run2.font.color.rgb = color
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    run3.font.name = "Calibri"
    run3.font.size = Pt(size)
    if color:
        run3.font.color.rgb = color
    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar_end)


if __name__ == "__main__":
    create_template("story-template.docx")
